import streamlit as st
from docx import Document as DocxDocument
from pdfminer.high_level import extract_text
import os
import base64
from langchain.chat_models import ChatOpenAI
from langchain.schema.messages import HumanMessage, AIMessage
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from PIL import Image
import io
import openai
import uuid  # Import UUID library for generating unique IDs

# Function to read docx file and process content
def read_docx(file_path):
    doc = DocxDocument(file_path)
    text_elements = []
    table_elements = []
    image_elements = []

    for para in doc.paragraphs:
        text_elements.append(para.text)

    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + "\t"
            table_text += "\n"
        table_elements.append(table_text)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_elements.append(rel.target_part.blob)

    return text_elements, table_elements, image_elements

# Function to read pdf file and process content
def read_pdf(file_path):
    text = extract_text(file_path)
    # For tables and images extraction, you might need additional libraries like pdfplumber or PyMuPDF
    # Extracting tables and images from PDF is more complex and depends on the PDF structure and content
    table_elements = []  # Placeholder for tables (needs implementation)
    image_elements = []  # Placeholder for images (needs implementation)
    return [text], table_elements, image_elements  # Returning text as a list for consistency with read_docx

# Function to encode image to base64
def encode_image(image_blob):
    return base64.b64encode(image_blob).decode('utf-8')

# Function to decode base64 to image
def decode_image(encoded_image):
    return Image.open(io.BytesIO(base64.b64decode(encoded_image)))

# Function to summarize image using GPT-4 Vision and cache results
@st.cache_data(show_spinner=False)
def summarize_image(encoded_image, _chain_gpt_4_vision):
    prompt = [
        AIMessage(content="You are a bot that is good at analyzing images."),
        HumanMessage(content=[
            {"type": "text", "text": "Describe the contents of this image."},
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{encoded_image}"
                },
            },
        ])
    ]
    response = _chain_gpt_4_vision.invoke(prompt)
    return response.content

# Function to get user input for file path using file uploader
def get_file_path():
    uploaded_file = st.file_uploader("Upload a file (DOCX or PDF)", type=["docx", "pdf"], key="file-uploader")
    if uploaded_file is not None:
        file_type = uploaded_file.type
        if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':  # Check if DOCX
            file_path = "temp.docx"
        elif file_type == 'application/pdf':  # Check if PDF
            file_path = "temp.pdf"
        else:
            st.error("Unsupported file type. Please upload a DOCX or PDF file.")
            return None

        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        return file_path
    else:
        return None

# Function to find relevant content based on query and threshold
def find_relevant_content(query, threshold, model, image_embeddings, image_summaries, image_elements):
    query_embedding = model.encode([query], convert_to_tensor=True)
    similarity_to_query = cosine_similarity(query_embedding, image_embeddings)
    highest_similarity_score = np.max(similarity_to_query)
    relevant_image_idx = np.argmax(similarity_to_query)

    if highest_similarity_score < threshold:
        return None, None

    relevant_image_summary = image_summaries[relevant_image_idx]
    relevant_image_blob = image_elements[relevant_image_idx]
    return relevant_image_summary, relevant_image_blob

# Function to query GPT for relevant information
def query_gpt(query, relevant_texts):
    prompt = "Based on the following information:\n\n"
    for text in relevant_texts:
        prompt += f"{text}\n\n"
    prompt += f"Answer the following query: {query}"
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500
    )
    return response.choices[0].message['content'].strip()

# Function to explain the image summary
def explain_image_summary(image_summary):
    prompt = f"""
    Explain the following summary without using the word 'image':
    '{image_summary}'
    To what category does it belong? Explain a bit about the category.
    Note: Any category should not be photography or art.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500
    )
    return response.choices[0].message['content'].strip()

# Function to generate a unique key
def get_unique_key(base_key):
    return str(uuid.uuid4())[:8] + "_" + base_key

# Main function
def main():
    st.title("Document Analysis")

    # Sidebar elements
    st.sidebar.markdown("<h1 style='color: white; font-size: 24px; font-weight: bold;'>Document Analyzer</h1>", unsafe_allow_html=True)
    st.sidebar.header("Controls:")
    
    # API key management
    if 'api_key_entered' not in st.session_state:
        st.session_state.api_key_entered = False

    if st.session_state.api_key_entered:
        if st.sidebar.button("Change or Edit Key"):
            st.session_state.api_key_entered = False
    else:
        key = st.sidebar.text_input("Please enter your OpenAI API key:", type="password")
        if st.sidebar.button("Submit Key"):
            openai.api_key = key
            os.environ['OPENAI_API_KEY'] = key
            st.session_state.api_key_entered = True
            st.experimental_rerun()

    # Reset button
    if st.sidebar.button("Reset Chat"):
        st.session_state.old_chats = []

    # Initialize old_chats if it doesn't exist
    if 'old_chats' not in st.session_state:
        st.session_state.old_chats = []

    # Show the remaining controls only if API key is entered
    if st.session_state.api_key_entered:
        default_threshold = 0.365
        
        # Generate unique keys for threshold input
        threshold_key = get_unique_key("threshold-input")
        
        # Text input for threshold value with dynamic key
        threshold_input = st.sidebar.text_input("Enter the threshold value (0-1):", default_threshold, key=threshold_key)
        
        # Validate the input and update threshold if valid
        try:
            threshold = float(threshold_input)
            if threshold < 0 or threshold > 1:
                st.sidebar.warning("Threshold must be between 0 and 1.")
                st.session_state.threshold_valid = False
            else:
                st.session_state.threshold_valid = True
        except ValueError:
            st.sidebar.warning("Please enter a numerical value.")
            st.session_state.threshold_valid = False

        if st.session_state.threshold_valid:
            st.sidebar.slider("Threshold Slider", min_value=0.0, max_value=1.0, step=0.001, value=threshold, key="threshold-slider")
        
        if 'show_summaries' not in st.session_state:
            st.session_state.show_summaries = False

        if st.sidebar.button("Show/Hide Summaries"):
            st.session_state.show_summaries = not st.session_state.show_summaries

        # Display the summaries if the button is toggled to show
        if st.session_state.show_summaries:
            if 'image_summaries' in st.session_state:
                st.sidebar.write("Image Summaries:")
                for i, summary in enumerate(st.session_state.image_summaries):
                    st.sidebar.write(f"Summary of image {i + 1}: {summary}")
            else:
                st.sidebar.write("No image summaries available.")
        
        # Read the document file
        file_path = get_file_path()
        if file_path is not None:
            with st.spinner("Processing file..."):
                if file_path.endswith('.docx'):
                    text_elements, table_elements, image_elements = read_docx(file_path)
                elif file_path.endswith('.pdf'):
                    text_elements, table_elements, image_elements = read_pdf(file_path)
                else:
                    st.error("Unsupported file type. Please upload a DOCX or PDF file.")
                    return

                # Initialize summarizers
                chain_gpt_4_vision = ChatOpenAI(model="gpt-4-vision-preview", max_tokens=1024)

                # Summarize images
                encoded_images = [encode_image(img) for img in image_elements]
                image_summaries = [summarize_image(ei, chain_gpt_4_vision) for ei in encoded_images]

                # Cache image summaries
                st.session_state.image_summaries = image_summaries

                # Load sentence transformer model for embeddings
                model = SentenceTransformer('all-MiniLM-L6-v2')

                # Generate image summary embeddings
                image_embeddings = model.encode(image_summaries, convert_to_tensor=True)

            # Display old chats at the top of the page
            st.write("Chat History:")
            if 'old_chats' in st.session_state:
                for idx, chat in enumerate(st.session_state.old_chats):
                    query, relevant_image_summary, relevant_image, explanation, answer = chat
                    st.markdown(f"**<span style='font-size:20px;'>Query {idx + 1}:</span>** {query}", unsafe_allow_html=True)
                    st.markdown(f"**Relevant Image Summary {idx + 1}:** {relevant_image_summary}")
                    st.image(relevant_image, caption=f"Image {idx + 1}")
                    st.markdown(f"**Explanation {idx + 1}:** {explanation}")
                    st.markdown(f"**Answer {idx + 1}:** {answer}")

            # Placeholder for results
            results_placeholder = st.empty()

            # Query section
            query_input = st.text_input("Enter your query:")

            # Submit query logic
            if st.button("Submit") or st.session_state.get('query_submit', False):
                if query_input:
                    st.session_state.query_submit = False
                    with results_placeholder.container():
                        relevant_image_summary, relevant_image_blob = find_relevant_content(query_input, threshold, model, image_embeddings, image_summaries, image_elements)

                        if relevant_image_summary is None:
                            st.write("No matching found")
                        else:
                            # Display relevant image summary
                            st.markdown(f"**Relevant Image Summary:** {relevant_image_summary}")

                            # Decode and display image
                            relevant_image = decode_image(encode_image(relevant_image_blob))
                            st.image(relevant_image, caption="Relevant Image")

                            # Explain the image summary
                            explanation = explain_image_summary(relevant_image_summary)
                            st.markdown(f"**Explanation:** {explanation}")

                            # Query GPT for an answer based on the document content
                            answer = query_gpt(query_input, text_elements + table_elements)
                            st.markdown(f"**Answer:** {answer}")

                            # Save chat
                            st.session_state.old_chats.append((query_input, relevant_image_summary, relevant_image, explanation, answer))

                            # Scroll to the bottom of the page
                            st.markdown('<style>div.css-1l02zno{height:80vh;}</style>', unsafe_allow_html=True)

            # JavaScript to handle enter key press
            st.markdown("""
            <script>
                document.addEventListener('DOMContentLoaded', function() {
                    const input = document.querySelector('.stTextInput > input');
                    input.addEventListener('keypress', function(e) {
                        if (e.key === 'Enter') {
                            var enter_pressed = true;
                            var query_submit = true;
                            Streamlit.setComponentValue({'enter_pressed': enter_pressed, 'query_submit': query_submit});
                        }
                    });
                });
            </script>
            """, unsafe_allow_html=True)           

             # Scroll to bottom button in the sidebar
            if st.sidebar.button("Scroll to Bottom"):
                st.markdown("""<script>window.scrollTo(0, document.body.scrollHeight);</script>""", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
